"""
render_docx.py — BundleModel -> DOCX bytes (engine/, no I/O beyond building the doc).

Walks the render-agnostic model produced by bundle_spec. Adding a section KIND here is a
render concern only; the spec decides WHAT is in a bundle, the renderer decides how it
looks. Editable DOCX is the internal companion to the authoritative PDF.
"""
from __future__ import annotations

import io
from typing import Any, Dict, List

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

NAVY = RGBColor(0x15, 0x65, 0xC0)
SLATE = RGBColor(0x37, 0x47, 0x4F)
GREY = RGBColor(0x6B, 0x72, 0x80)
RED = RGBColor(0xB7, 0x1C, 0x1C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _cell_bg(cell, hex_color: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcpr.append(shd)


def _kv(doc, label: str, value: str, color=None) -> None:
    p = doc.add_paragraph()
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = SLATE
    v = p.add_run(str(value))
    v.font.size = Pt(10)
    if color:
        v.font.color.rgb = color


def _h(doc, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = NAVY if level == 1 else SLATE


def _table(doc, headers: List[str], rows: List[List[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, hcell in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = hcell
        for r in c.paragraphs[0].runs:
            r.bold = True
            r.font.color.rgb = WHITE
        _cell_bg(c, "1565C0")
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = str(val)
        if ri % 2 == 0:
            for c in t.rows[ri].cells:
                _cell_bg(c, "EEF2F6")


# ── per-kind renderers ───────────────────────────────────────────────────────

def _r_cover(doc, sec, meta):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(sec["title"])
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = NAVY
    ps = doc.add_paragraph()
    ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = ps.add_run(sec.get("subtitle", ""))
    rs.font.size = Pt(12)
    rs.font.color.rgb = SLATE
    pc = doc.add_paragraph()
    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = pc.add_run(sec.get("city", ""))
    rc.font.size = Pt(16)
    rc.bold = True
    rc.font.color.rgb = SLATE
    for _ in range(2):
        doc.add_paragraph()
    for label, key in [("Prepared for", "audience")]:
        pp = doc.add_paragraph()
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = pp.add_run(f"{label}: {sec.get(key, '')}")
        rr.font.size = Pt(11)
        rr.font.color.rgb = GREY
    # Document control block — tamper-evidence
    dc = doc.add_paragraph()
    dc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lines = (f"Document ID: {meta['doc_id']}   |   Generated: {meta['generated_utc']}\n"
             f"Tool release: {meta['tool_release']}   |   "
             f"Content SHA-256: {meta['content_sha256']}")
    rdc = dc.add_run(lines)
    rdc.font.size = Pt(8)
    rdc.font.color.rgb = GREY
    doc.add_page_break()


def _r_provenance(doc, sec, meta):
    _h(doc, sec["title"], 1)
    p = doc.add_paragraph()
    r = p.add_run(sec["headline"])
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = SLATE
    doc.add_paragraph()
    rows = [[row["source"], str(row["count"]),
             "Discovered" if row["discovered"] else "Declared"] for row in sec["rows"]]
    if rows:
        _table(doc, ["Source", "Count", "Origin"], rows)


def _r_exec(doc, sec, meta):
    _h(doc, sec["title"], 1)
    for f in sec["facts"]:
        _kv(doc, f["label"], f["value"])
    doc.add_paragraph()
    doc.add_paragraph(sec["background"])
    n = doc.add_paragraph()
    rn = n.add_run(sec["notice"])
    rn.italic = True
    rn.font.size = Pt(9)
    rn.font.color.rgb = GREY


def _r_inventory(doc, sec, meta):
    _h(doc, sec["title"], 1)
    doc.add_paragraph(sec["intro"])
    doc.add_paragraph()
    rows = sec.get("rows", [])
    if not rows:
        doc.add_paragraph(sec.get("empty", "None."))
        return
    _table(doc, ["Vendor / System", "Type", "Source", "Confidence", "Status"],
           [[r["vendor"], r["type"], r["source"], r["confidence"], r["status"]] for r in rows])


def _r_detail(doc, sec, meta):
    _h(doc, sec["title"], 1)
    for f in sec["facts"]:
        _kv(doc, f["label"], f["value"])
    doc.add_paragraph()
    doc.add_paragraph(sec["note"])


def _r_violations(doc, sec, meta):
    _h(doc, sec["title"], 1)
    items = sec.get("items", [])
    if not items:
        doc.add_paragraph(sec.get("empty", "None."))
        return
    doc.add_paragraph(sec["intro"])
    for i, v in enumerate(items, start=1):
        _h(doc, f"Finding {i}: {v['rule_id']} — {v['citation']}", 2)
        _kv(doc, "Severity", v["severity"], RED if v["severity"] == "HIGH" else None)
        _kv(doc, "Status", v["status"])
        _kv(doc, "First observed", v["first_observed"])
        _kv(doc, "Cure deadline", v["cure_deadline"])
        if v.get("days_remaining") is not None:
            _kv(doc, "Days remaining", str(v["days_remaining"]))
        if v.get("page_url"):
            _kv(doc, "Location", v["page_url"])
        if v.get("indicators"):
            _kv(doc, "Matched indicators", ", ".join(str(x) for x in v["indicators"]))
        if v.get("remediation"):
            _kv(doc, "Recommended remediation", v["remediation"])
        doc.add_paragraph()


def _r_recommendations(doc, sec, meta):
    _h(doc, sec["title"], 1)
    for item in sec["items"]:
        doc.add_paragraph(item, style="List Bullet")


def _r_statutory(doc, sec, meta):
    _h(doc, sec["title"], 1)
    for it in sec["items"]:
        _kv(doc, it["citation"], it["text"])


def _r_methodology(doc, sec, meta):
    _h(doc, sec["title"], 1)
    for para in sec["paragraphs"]:
        doc.add_paragraph(para)


def _r_attestation(doc, sec, meta):
    _h(doc, sec["title"], 1)
    doc.add_paragraph(sec["statement"])
    doc.add_paragraph()
    for field in sec["fields"]:
        p = doc.add_paragraph()
        r = p.add_run(f"{field}: ")
        r.bold = True
        p.add_run("____________________________")


_RENDERERS = {
    "cover": _r_cover, "provenance_summary": _r_provenance, "exec_status": _r_exec,
    "asset_inventory": _r_inventory, "compliance_detail": _r_detail,
    "violations": _r_violations, "recommendations": _r_recommendations,
    "statutory_reference": _r_statutory, "methodology": _r_methodology,
    "attestation": _r_attestation,
}


def render(model: Dict[str, Any]) -> bytes:
    """BundleModel -> .docx bytes."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)
    meta = model["meta"]
    for sec in model["sections"]:
        fn = _RENDERERS.get(sec["kind"])
        if fn:
            fn(doc, sec, meta)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
