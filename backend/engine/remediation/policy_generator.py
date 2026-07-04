"""
policy_generator.py -- Vendor-specific AI Use Policy generator for TRAIGA compliance.

Generates a Word (.docx) AI Use Policy document customised to the AI vendors
actually detected on a city's public website. Callable from the remediation
API route and (future) CLI.

Compliance basis: Texas HB 149 / TRAIGA, Tex. Bus. & Com. Code Ch. 552
Effective: January 1, 2026
"""
from __future__ import annotations

import json
import os
import tempfile
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError as _e:
    raise ImportError("python-docx not installed. Run: pip install python-docx") from _e

# ---------------------------------------------------------------------------
# Vendor intelligence profiles
# ---------------------------------------------------------------------------
VENDOR_PROFILES: Dict[str, Dict[str, Any]] = {
    "citibot": {
        "display_name":   "Citibot",
        "description":    (
            "AI-powered virtual assistant providing automated responses to 311 "
            "inquiries, frequently asked questions, and resident service requests "
            "via the City website and mobile channels."
        ),
        "primary_use":    "Resident-facing virtual assistant / chatbot",
        "data_processed": "Resident inquiries, service request submissions, interaction logs",
        "biometric_risk": False,
        "vendor_url":     "https://www.citibot.io",
    },
    "civicplus": {
        "display_name":   "CivicPlus",
        "description":    (
            "Municipal website content management platform with integrated AI-assisted "
            "citizen engagement and conversational interface features."
        ),
        "primary_use":    "Municipal website platform with AI chat features",
        "data_processed": "Resident interactions, form submissions, service requests",
        "biometric_risk": False,
        "vendor_url":     "https://www.civicplus.com",
    },
    "govpilot": {
        "display_name":   "GovPilot",
        "description":    (
            "Cloud-based government management platform incorporating AI-driven "
            "process automation and resident-facing chatbot capabilities."
        ),
        "primary_use":    "Government process automation with AI components",
        "data_processed": "Permit applications, resident records, workflow data",
        "biometric_risk": False,
        "vendor_url":     "https://www.govpilot.com",
    },
    "municode": {
        "display_name":   "Municode",
        "description":    (
            "Municipal code publishing platform with AI-assisted search and "
            "document navigation for public ordinances and regulations."
        ),
        "primary_use":    "AI-assisted municipal code search and navigation",
        "data_processed": "Search queries, document interaction logs",
        "biometric_risk": False,
        "vendor_url":     "https://www.municode.com",
    },
}

_PRIMARY   = RGBColor(0x1A, 0x23, 0x7E)
_SECONDARY = RGBColor(0x28, 0x3B, 0x93)
_GRAY      = RGBColor(0x60, 0x60, 0x60)


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = _PRIMARY if level == 1 else _SECONDARY
    return p


def _bold_para(doc: Document, text: str):
    p = doc.add_paragraph()
    p.add_run(text).bold = True
    return p


def _styled_para(doc: Document, text: str, italic: bool = False,
                 color: Optional[RGBColor] = None):
    p = doc.add_paragraph(text)
    for run in p.runs:
        if italic:
            run.italic = True
        if color:
            run.font.color.rgb = color
    return p


def _add_hr(doc: Document) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1A237E")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _numbered_list(doc: Document, items: List[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def _bullet_list(doc: Document, items: List[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _vendor_table(doc: Document, assets: List[Dict[str, Any]]) -> None:
    headers = ["System Name", "Primary Use", "Data Processed", "Biometric Risk", "Vendor URL"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Shading Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
    for asset in assets:
        vid  = asset.get("vendor_id", "").lower()
        prof = VENDOR_PROFILES.get(vid, {})
        row  = table.add_row().cells
        row[0].text = prof.get("display_name") or asset.get("display_name") or vid.title()
        row[1].text = prof.get("primary_use", "AI-assisted municipal service")
        row[2].text = prof.get("data_processed", "Resident interaction data")
        row[3].text = "Yes" if prof.get("biometric_risk") else "No"
        row[4].text = prof.get("vendor_url", "N/A")


def _violation_summary(violations: List[Dict]) -> str:
    if not violations:
        return "No open violations identified at time of policy adoption."
    rule_ids = sorted({v.get("rule_id", "unknown") for v in violations})
    return (
        "As of the scan date, {} open violation(s) were identified "
        "under the following rule categories: {}. "
        "This policy is adopted as part of the cure period remediation plan."
    ).format(len(violations), ", ".join(rule_ids))


def _vendor_disclosure_para(assets: List[Dict[str, Any]]) -> str:
    if not assets:
        return (
            "The City currently does not deploy externally observable AI systems "
            "on its public website. Should AI systems be deployed in the future, "
            "this policy requires disclosure prior to public use."
        )
    names = []
    for a in assets:
        vid  = a.get("vendor_id", "").lower()
        prof = VENDOR_PROFILES.get(vid, {})
        names.append(prof.get("display_name") or a.get("display_name") or vid.title())
    if len(names) > 1:
        listed = ", ".join(names[:-1]) + " and " + names[-1]
    else:
        listed = names[0]
    return (
        "The City currently deploys the following AI system(s) on its public-facing "
        "digital properties: {}. Each system is subject to the disclosure, "
        "transparency, and data governance requirements set forth in this policy "
        "and in Tex. Bus. & Com. Code Ch. 552."
    ).format(listed)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def generate_ai_use_policy(
    city: str,
    jurisdiction: str,
    domain: str,
    detected_assets: List[Dict[str, Any]],
    violations: List[Dict[str, Any]],
    output_path: Optional[str] = None,
    scan_date: Optional[str] = None,
) -> str:
    """
    Generate a vendor-specific AI Use Policy Word document for a municipality.

    Returns the absolute path to the generated .docx file.
    """
    if output_path is None:
        tmp = tempfile.mkdtemp()
        safe = city.replace(" ", "_").replace("/", "_")
        output_path = os.path.join(tmp, safe + "_AI_Use_Policy.docx")

    now_iso = scan_date or datetime.now(timezone.utc).isoformat()
    try:
        scan_dt = datetime.fromisoformat(now_iso.replace("Z", "+00:00"))
        adopted = scan_dt.strftime("%B %d, %Y")
    except ValueError:
        adopted = datetime.now(timezone.utc).strftime("%B %d, %Y")

    review_year = datetime.now(timezone.utc).year + 1
    policy_id   = "POL-AI-" + str(uuid.uuid4())[:6].upper()
    city_ref    = city + ' ("the City")'
    domain_ref  = domain or "the City website"

    doc = Document()

    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width  = Inches(8.5)
    section.left_margin = section.right_margin = Inches(1.25)
    section.top_margin  = section.bottom_margin = Inches(1.0)

    # Cover
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(city.upper())
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = _PRIMARY

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_r = sub_p.add_run("ARTIFICIAL INTELLIGENCE USE POLICY")
    sub_r.bold = True
    sub_r.font.size = Pt(14)
    sub_r.font.color.rgb = _SECONDARY

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.add_run(
        "Policy ID: {}  |  Adopted: {}  |  Review Date: December 31, {}".format(
            policy_id, adopted, review_year
        )
    ).font.color.rgb = _GRAY

    basis_p = doc.add_paragraph()
    basis_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    basis_r = basis_p.add_run(
        "Compliance Basis: Texas HB 149 (TRAIGA)  |  "
        "Tex. Bus. & Com. Code Ch. 552  |  Effective January 1, 2026"
    )
    basis_r.italic = True
    basis_r.font.color.rgb = _GRAY

    _add_hr(doc)
    doc.add_paragraph()

    # Section 1
    _heading(doc, "Section 1 -- Purpose and Legal Authority")
    doc.add_paragraph(
        "This policy establishes the {} framework for the responsible, transparent, "
        "and lawful use of Artificial Intelligence (AI) systems in public-facing "
        "municipal services. It is adopted in compliance with the Texas Responsible "
        "Artificial Intelligence Governance Act (TRAIGA), enacted as Texas HB 149, "
        "89th Legislature, Regular Session, and codified at Tex. Bus. & Com. Code, "
        "Chapter 552, with an effective date of January 1, 2026.".format(city_ref)
    )
    doc.add_paragraph(
        "The City recognises that AI systems deployed in resident-facing services "
        "must be operated transparently, with clear disclosure to the public, "
        "appropriate data governance, and defined accountability mechanisms consistent "
        "with the statutory obligations imposed by Chapter 552."
    )

    # Section 2
    _heading(doc, "Section 2 -- Scope and Applicability")
    doc.add_paragraph(
        "This policy applies to all AI systems deployed by the City or on behalf "
        "of the City on its public-facing digital properties, including {} "
        "and any affiliated portals, mobile applications, or chatbot interfaces "
        "accessible to residents. It applies to all City departments, vendors, "
        "and third-party service providers operating AI systems under contract "
        "with the City.".format(domain_ref)
    )
    doc.add_paragraph(
        "This policy does not apply to AI systems used exclusively for internal "
        "City operations that do not process resident-facing interactions or "
        "publicly observable outputs. Separate governance documentation governs "
        "internal AI use."
    )

    # Section 3
    _heading(doc, "Section 3 -- AI System Inventory")
    doc.add_paragraph(
        "The following AI systems have been identified on the City's public-facing "
        "digital properties as of the most recent external compliance scan. This "
        "inventory must be reviewed and updated no less than annually, and within "
        "30 days of any new AI system deployment or decommission."
    )
    doc.add_paragraph()
    if detected_assets:
        _vendor_table(doc, detected_assets)
        doc.add_paragraph()
    else:
        _styled_para(
            doc,
            "No AI systems were detected on the City's public website at the time "
            "of the most recent scan. This inventory will be updated upon deployment "
            "of any AI-powered system.",
            italic=True,
        )

    # Section 4
    _heading(doc, "Section 4 -- Public Disclosure Requirements (Tex. Bus. & Com. Code ss.552.051)")
    doc.add_paragraph(
        "Chapter 552 requires governmental entities that deploy AI systems in "
        "resident interactions to provide conspicuous, plain-language disclosure "
        "that the resident is interacting with an AI system, not a human employee."
    )
    doc.add_paragraph(_vendor_disclosure_para(detected_assets))
    _bold_para(doc, "Required Disclosure Standards:")
    _bullet_list(doc, [
        "Disclosure must appear at or before the initiation of any AI-mediated "
        "interaction (e.g., at the start of a chatbot session or adjacent to an "
        "AI-generated content block).",
        "Disclosure language must be in plain English and legible to a general audience.",
        "The name or category of the AI system must be identified.",
        (
            "Example compliant language: \"You are interacting with an AI-powered "
            "virtual assistant. This system is not a City employee. For human "
            "assistance, please contact [department contact].\""
        ),
        "Disclosure must be maintained current and updated within 30 days of "
        "any change in AI system deployment.",
    ])
    for asset in detected_assets:
        vid  = asset.get("vendor_id", "").lower()
        prof = VENDOR_PROFILES.get(vid, {})
        name = prof.get("display_name") or asset.get("display_name") or vid.title()
        desc = prof.get("description", "AI-assisted resident service system.")
        _bold_para(doc, name + " -- Required Disclosure:")
        doc.add_paragraph(
            "{} The City shall ensure that all resident-facing interfaces powered by "
            "{} display the required TRAIGA disclosure statement prior to or at the "
            "commencement of any AI-mediated interaction.".format(desc, name)
        )

    # Section 5
    _heading(doc, "Section 5 -- Biometric Data Protections (Tex. Bus. & Com. Code ss.552.052)")
    biometric_assets = [
        a for a in detected_assets
        if VENDOR_PROFILES.get(a.get("vendor_id", "").lower(), {}).get("biometric_risk")
    ]
    if biometric_assets:
        bio_names = ", ".join(
            VENDOR_PROFILES.get(a.get("vendor_id", "").lower(), {}).get("display_name", "")
            for a in biometric_assets
        )
        doc.add_paragraph(
            "The following AI systems deployed by the City have been flagged for "
            "potential biometric data processing: {}. The City shall comply with all "
            "requirements of ss.552.052 regarding collection, storage, and use of "
            "biometric identifiers.".format(bio_names)
        )
    else:
        doc.add_paragraph(
            "No AI systems currently deployed by the City have been identified as "
            "processing biometric data. Should any future system require biometric "
            "data collection (facial recognition, voiceprint, fingerprint, or other "
            "biometric identifiers), the City shall comply with all requirements of "
            "Tex. Bus. & Com. Code ss.552.052 prior to deployment, including resident "
            "consent, data retention limits, and prohibition on sale of biometric data."
        )
    _bullet_list(doc, [
        "Biometric data may not be sold, leased, traded, or otherwise profited from.",
        "Retention schedules for any biometric data must be established in writing "
        "prior to collection.",
        "Residents must provide informed written consent prior to biometric data "
        "collection by any City AI system.",
    ])

    # Section 6
    _heading(doc, "Section 6 -- Resident Rights and Transparency")
    doc.add_paragraph(
        "The City affirms the following rights for all residents interacting with "
        "City AI systems:"
    )
    _numbered_list(doc, [
        "The right to be informed when interacting with an AI system rather than "
        "a human City employee.",
        "The right to request human review of any AI-generated decision, "
        "recommendation, or determination that affects the resident.",
        "The right to opt out of AI-mediated interactions and receive equivalent "
        "service through a human channel.",
        "The right to request information about what data is processed by City "
        "AI systems and how it is used.",
        "The right to file a complaint regarding AI system interactions through "
        "the City's established grievance process.",
    ])

    # Section 7
    _heading(doc, "Section 7 -- Cure Period Compliance (Tex. Bus. & Com. Code ss.552.053)")
    doc.add_paragraph(
        "Chapter 552 provides a 60-day cure period following identification of "
        "a violation, during which the City may remedy the non-compliance without "
        "penalty. This policy is adopted as part of the City's good-faith cure "
        "period remediation plan."
    )
    doc.add_paragraph(_violation_summary(violations))
    _bold_para(doc, "Cure Period Action Items:")
    _numbered_list(doc, [
        "Within 15 days: City Attorney and IT leadership review this policy and "
        "confirm adoption by City Council or equivalent governing authority.",
        "Within 30 days: All AI disclosure language updated on City website and "
        "resident-facing interfaces to comply with ss.552.051.",
        "Within 45 days: Staff training completed for all departments deploying "
        "AI systems.",
        "Within 60 days: All open violations remediated and documented. "
        "Compliance scan re-run to confirm cure.",
        "Ongoing: Annual policy review conducted no later than December 31 each year.",
    ])

    # Section 8
    _heading(doc, "Section 8 -- Enforcement and Penalties (Tex. Bus. & Com. Code ss.552.054)")
    doc.add_paragraph(
        "Violations of Chapter 552 are subject to enforcement by the Texas Attorney "
        "General. The City designates the following accountability structure for AI governance:"
    )
    _bullet_list(doc, [
        "Chief Information Officer (CIO): Overall AI governance accountability, "
        "annual compliance review, and regulatory interface.",
        "City Attorney: Legal review of AI vendor contracts, disclosure language, "
        "and cure period documentation.",
        "Department Directors: Compliance within their respective departments; "
        "reporting of new AI system deployments within 10 business days.",
        "IT Leadership: Technical implementation of disclosure requirements, "
        "vendor management, and AI system inventory maintenance.",
    ])

    # Section 9
    _heading(doc, "Section 9 -- Policy Review and Update Schedule")
    doc.add_paragraph(
        "This policy shall be reviewed and updated no less than annually, with the "
        "next scheduled review no later than December 31, {}. "
        "Reviews shall also be triggered by: (a) deployment of a new AI system; "
        "(b) material changes to TRAIGA or related regulations; "
        "(c) identification of a new compliance violation; or "
        "(d) significant change in AI vendor contracts or capabilities.".format(review_year)
    )

    # Adoption block
    _add_hr(doc)
    doc.add_paragraph()
    _heading(doc, "Policy Adoption", level=2)
    doc.add_paragraph(
        "This AI Use Policy is adopted by {}, {}, "
        "effective upon approval by the governing authority.".format(city, jurisdiction)
    )
    doc.add_paragraph()
    for label in [
        "City Manager / Authorized Official:",
        "Date:",
        "City Attorney Review:",
        "Date:",
    ]:
        sig_p = doc.add_paragraph()
        sig_p.add_run(label + "   ").bold = True
        sig_p.add_run("_" * 40)

    doc.add_paragraph()
    _styled_para(
        doc,
        "DISCLAIMER: This document is a draft AI Use Policy generated by the "
        "AI Transparency Auditor Platform based on external compliance scan data. "
        "It is intended as a starting point for legal review and City adoption. "
        "It does not constitute legal advice. All findings should be reviewed by "
        "qualified legal counsel before official adoption.",
        italic=True,
        color=_GRAY,
    )

    doc.save(output_path)
    return output_path
