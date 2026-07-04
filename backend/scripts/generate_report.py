"""
generate_report.py -- TRAIGA Compliance Report Generator

Usage:
    python generate_report.py \
        --city "City Name" \
        --scorecard '{"traiga_status":"in_cure",...}' \
        --violations '[{...}]' \
        --output "/path/to/report.docx" \
        [--brand-config '{"company_name":"Acme","primary_color":"#1A237E"}']

Brand config keys (all optional):
    company_name, tagline, logo_path, primary_color,
    secondary_color, website, confidentiality
"""
from __future__ import annotations
import argparse, json, sys, uuid
from datetime import datetime, timezone
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError as _docx_err:

    raise ImportError("python-docx is not installed. Run: pip install python-docx --break-system-packages") from _docx_err

DEFAULT_BRAND = {
    "company_name":    "AI Transparency Auditor Platform",
    "tagline":         "Municipal AI Compliance -- Texas TRAIGA",
    "logo_path":       None,
    "primary_color":   "#1A237E",
    "secondary_color": "#283B93",
    "website":         None,
    "confidentiality": "CONFIDENTIAL -- FOR OFFICIAL USE ONLY",
}

def _load_brand(brand_json):
    brand = DEFAULT_BRAND.copy()
    if brand_json:
        try:
            brand.update({k: v for k, v in json.loads(brand_json).items() if v is not None})
        except json.JSONDecodeError as e:
            print(f"WARNING: Bad brand config ({e}), using defaults.")
    return brand

def _hex_to_rgb(h):
    h = h.lstrip("#")
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def _add_field(run, field_name):
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"),"begin")
    t = OxmlElement("w:instrText"); t.set(qn("xml:space"),"preserve"); t.text=f" {field_name} "
    s = OxmlElement("w:fldChar"); s.set(qn("w:fldCharType"),"separate")
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"),"end")
    for el in (b,t,s,e): run._r.append(el)

def _border(para, color_hex="1A237E", size=12):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"),"single"); top.set(qn("w:sz"),str(size))
    top.set(qn("w:space"),"1"); top.set(qn("w:color"),color_hex.lstrip("#"))
    pBdr.append(top); pPr.append(pBdr)

def _cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)

def _safe(d, *keys, default="N/A"):
    v = d
    for k in keys:
        if not isinstance(v, dict): return default
        v = v.get(k)
        if v is None: return default
    return str(v) if str(v) not in ("","None","NaN") else default


STATUTE_BASE = "https://statutes.capitol.texas.gov/Docs/BC/htm/BC.552.htm"
HB149_URL    = "https://capitol.texas.gov/tlodocs/89R/billtext/html/HB00149F.htm"

def _citation_url(citation_str):
    """Extract section anchor from a citation string and return the statute URL."""
    import re
    m = re.search(r"§(552\.\d+)", citation_str)
    if m:
        return f"{STATUTE_BASE}#{m.group(1)}"
    return STATUTE_BASE

def _add_hyperlink(para, text, url, color=None):
    """Add a clickable hyperlink run to an existing paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    # Add relationship
    part = para.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    # Build hyperlink element
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    hyperlink.set(qn("w:history"), "1")
    # Run inside hyperlink
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    # Underline
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    # Colour
    c = OxmlElement("w:color")
    hex_val = (color.lstrip("#") if color else "1A237E")
    c.set(qn("w:val"), hex_val); rPr.append(c)
    # Style (hyperlink char style if available)
    rStyle = OxmlElement("w:rStyle"); rStyle.set(qn("w:val"), "Hyperlink"); rPr.append(rStyle)
    new_run.append(rPr)
    t = OxmlElement("w:t"); t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    para._p.append(hyperlink)
    return para

SEVERITY_COLORS = {
    "high":   RGBColor(0xC6,0x28,0x28),
    "medium": RGBColor(0xF5,0x7F,0x17),
    "low":    RGBColor(0x15,0x65,0xC0),
}
STATUS_COLORS = {
    "compliant":     RGBColor(0x2E,0x7D,0x32),
    "in_cure":       RGBColor(0xF5,0x7F,0x17),
    "non_compliant": RGBColor(0xC6,0x28,0x28),
    "expired":       RGBColor(0x6A,0x1B,0x9A),
    "not_assessed":  RGBColor(0x42,0x42,0x42),
}
STATUS_LABELS = {
    "compliant":     "COMPLIANT",
    "in_cure":       "IN CURE PERIOD",
    "non_compliant": "NON-COMPLIANT",
    "expired":       "CURE PERIOD EXPIRED",
    "not_assessed":  "NOT YET ASSESSED",
}
STATUS_PLAIN = {
    "compliant":     "This municipality has met all current TRAIGA AI disclosure requirements.",
    "in_cure":       "One or more AI disclosure violations have been identified. The municipality is within the 60-day cure period established by Texas HB 149 and has the opportunity to remediate before enforcement action may be initiated.",
    "non_compliant": "This municipality has unresolved AI disclosure violations that require immediate attention. Failure to remediate may result in enforcement action under Tex. Bus. & Com. Code §552.",
    "expired":       "The 60-day cure period for one or more violations has expired without remediation. This municipality is at risk of enforcement action by the Texas Attorney General.",
    "not_assessed":  "This municipality has been registered as an audit target but has not yet been scanned. An audit should be initiated to establish compliance status.",
}
TRAIGA_PROVISIONS = [
    ("§552.051","Requirement to Disclose AI Use","A governmental entity that deploys or uses an artificial intelligence system that interacts with members of the public must conspicuously disclose the use of AI to each person who interacts with the system."),
    ("§552.052","Biometric Data","Additional consent and disclosure requirements apply when an AI system collects, stores, or processes biometric identifiers or biometric information."),
    ("§552.053","Cure Period","Upon identification of a violation, a governmental entity has 60 calendar days to cure the violation before the Attorney General may take enforcement action."),
    ("§552.054","Enforcement","The Attorney General may bring an action to enjoin a violation and seek civil penalties of up to $10,000 per violation per day."),
]

def _heading(doc, text, level, brand):
    p = doc.add_heading(text, level=level)
    color = _hex_to_rgb(brand["primary_color"] if level==1 else brand["secondary_color"])
    for run in p.runs: run.font.color.rgb = color

def _kv(doc, label, value, color=None, url=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    lb = p.add_run(f"{label}: "); lb.bold = True
    if url:
        _add_hyperlink(p, value, url, color.lstrip("#") if hasattr(color,"lstrip") else None)
    else:
        vr = p.add_run(value)
        if color: vr.font.color.rgb = color

def _setup_document(doc, brand):
    phex = brand["primary_color"].lstrip("#")
    prgb = _hex_to_rgb(brand["primary_color"])
    grey = RGBColor(0x75,0x75,0x75)
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)
        section.different_first_page_header_footer = True
        # Interior header
        hdr = section.header; hdr.is_linked_to_previous = False
        hp = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
        hp.clear(); hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.paragraph_format.space_after = Pt(4)
        _border(hp, phex, size=18)
        cr = hp.add_run(brand["company_name"]); cr.font.size=Pt(8); cr.font.color.rgb=prgb; cr.font.bold=True
        # Interior footer
        ftr = section.footer; ftr.is_linked_to_previous = False
        fp = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
        fp.clear(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(4)
        _border(fp, phex, size=6)
        r1 = fp.add_run(f"{brand['confidentiality']}  |  {brand['company_name']}  |  Page ")
        r1.font.size=Pt(8); r1.font.color.rgb=grey
        pr = fp.add_run(); pr.font.size=Pt(8); pr.font.color.rgb=grey; _add_field(pr,"PAGE")
        r2 = fp.add_run(" of "); r2.font.size=Pt(8); r2.font.color.rgb=grey
        nr = fp.add_run(); nr.font.size=Pt(8); nr.font.color.rgb=grey; _add_field(nr,"NUMPAGES")
        if brand.get("website"):
            wr = fp.add_run(f"  |  {brand['website']}"); wr.font.size=Pt(8); wr.font.color.rgb=grey
        # First-page footer (cover) -- confidentiality only
        fftr = section.first_page_footer; fftr.is_linked_to_previous = False
        ffp = fftr.paragraphs[0] if fftr.paragraphs else fftr.add_paragraph()
        ffp.clear(); ffp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = ffp.add_run(brand["confidentiality"]); sr.font.size=Pt(9); sr.font.color.rgb=grey; sr.italic=True

def _cover_page(doc, city, report_date, doc_id, brand):
    prgb = _hex_to_rgb(brand["primary_color"])
    srgb = _hex_to_rgb(brand["secondary_color"])
    grey = RGBColor(0x75,0x75,0x75)
    # Top accent rule
    rule = doc.add_paragraph(); rule.paragraph_format.space_before=Pt(0); rule.paragraph_format.space_after=Pt(0)
    _border(rule, brand["primary_color"].lstrip("#"), size=36)
    doc.add_paragraph(); doc.add_paragraph()
    # Optional logo
    if brand.get("logo_path") and Path(brand["logo_path"]).exists():
        lp = doc.add_paragraph(); lp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        lp.add_run().add_picture(brand["logo_path"], width=Inches(1.5))
        doc.add_paragraph()
    tp = doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run("AI TRANSPARENCY COMPLIANCE REPORT"); tr.bold=True; tr.font.size=Pt(22); tr.font.color.rgb=prgb
    sp = doc.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run("Texas HB 149 / TRAIGA Assessment"); sr.font.size=Pt(13); sr.font.color.rgb=srgb
    doc.add_paragraph()
    cp = doc.add_paragraph(); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run(city); cr.bold=True; cr.font.size=Pt(20)
    doc.add_paragraph()
    for label, value in [("Report Date",report_date),("Document ID",doc_id),("Authority","Tex. Bus. & Com. Code Ch. 552 | 89th Texas Legislature"),("Version","1.0")]:
        mp = doc.add_paragraph(); mp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        lb = mp.add_run(f"{label}: "); lb.bold=True; lb.font.size=Pt(10); lb.font.color.rgb=grey
        vl = mp.add_run(value); vl.font.size=Pt(10); vl.font.color.rgb=grey
    # Authority line with hyperlink to bill text
    auth_p = doc.add_paragraph(); auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grey = RGBColor(0x75,0x75,0x75)
    ar = auth_p.add_run("Authority: "); ar.bold=True; ar.font.size=Pt(10); ar.font.color.rgb=grey
    _add_hyperlink(auth_p, "Tex. Bus. & Com. Code Ch. 552 | Texas HB 149 (89th Legislature)", HB149_URL, "1A237E")
    doc.add_paragraph()
    cred = doc.add_paragraph(); cred.alignment=WD_ALIGN_PARAGRAPH.CENTER
    cr2 = cred.add_run(f"Prepared by {brand['company_name']}"); cr2.italic=True; cr2.font.size=Pt(10); cr2.font.color.rgb=grey
    if brand.get("tagline"):
        tag = doc.add_paragraph(); tag.alignment=WD_ALIGN_PARAGRAPH.CENTER
        tr2 = tag.add_run(brand["tagline"]); tr2.font.size=Pt(9); tr2.font.color.rgb=grey
    if brand.get("website"):
        wp = doc.add_paragraph(); wp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        wr = wp.add_run(brand["website"]); wr.font.size=Pt(9); wr.font.color.rgb=grey
    doc.add_paragraph()
    bottom = doc.add_paragraph(); _border(bottom, brand["primary_color"].lstrip("#"), size=12)
    doc.add_page_break()

def _executive_summary(doc, city, scorecard, violations, brand):
    _heading(doc, "Executive Summary", 1, brand)
    status = scorecard.get("traiga_status","not_assessed")
    label  = STATUS_LABELS.get(status, status.upper())
    scolor = STATUS_COLORS.get(status, STATUS_COLORS["not_assessed"])
    bp = doc.add_paragraph()
    br = bp.add_run(label); br.bold=True; br.font.size=Pt(14); br.font.color.rgb=scolor
    doc.add_paragraph(STATUS_PLAIN.get(status,""))
    doc.add_paragraph()
    _kv(doc,"Municipality",city)
    _kv(doc,"Compliance Score",f"{_safe(scorecard,'compliance_score')}/100", scolor)
    _kv(doc,"Open Violations",str(scorecard.get("open_violations_count", len([v for v in violations if v.get("status")=="in_cure"]))))
    _kv(doc,"Last Scanned",_safe(scorecard,"last_scanned_utc").split("T")[0])
    urgent = sorted([v for v in violations if v.get("days_remaining") is not None], key=lambda v: v.get("days_remaining") or 9999)
    if urgent:
        days = urgent[0].get("days_remaining","N/A")
        deadline = _safe(urgent[0],"cure_deadline_utc").split("T")[0]
        ucol = SEVERITY_COLORS["high"] if isinstance(days,int) and days<15 else SEVERITY_COLORS["medium"]
        _kv(doc,"Most Urgent Deadline",f"{days} days remaining (due {deadline})", ucol)
    doc.add_paragraph()
    bg = doc.add_paragraph(); bg.add_run("Background").bold=True
    doc.add_paragraph("Texas House Bill 149, codified as Tex. Bus. & Com. Code Ch. 552 (TRAIGA), requires Texas governmental entities that deploy AI systems interacting with the public to conspicuously disclose the use of those systems. Upon identification of a violation, a 60-day cure period is provided before enforcement action may be initiated by the Texas Attorney General.")

def _asset_inventory(doc, scorecard, brand):
    _heading(doc, "AI Asset Inventory", 1, brand)
    doc.add_paragraph("The following AI systems were detected on the municipality's public-facing website during the most recent automated scan. Detection uses behavioral fingerprinting and script-host analysis.")
    doc.add_paragraph()
    assets = scorecard.get("ai_assets",[])
    if not assets:
        doc.add_paragraph("No AI assets detected during the most recent scan.")
        return
    phex   = brand["primary_color"].lstrip("#")
    headers = ["Vendor / System","Asset Type","Page URL","Confidence","Status"]
    table = doc.add_table(rows=1+len(assets), cols=len(headers))
    table.style="Table Grid"; table.alignment=WD_TABLE_ALIGNMENT.LEFT
    for i,h in enumerate(headers):
        cell=table.rows[0].cells[i]; cell.text=h
        for run in cell.paragraphs[0].runs: run.bold=True; run.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        _cell_bg(cell, phex)
    for ri,asset in enumerate(assets,start=1):
        row=table.rows[ri]
        vals=[
            _safe(asset,"display_name"),
            (", ".join(asset["asset_type"]) if isinstance(asset.get("asset_type"),list) else _safe(asset,"asset_type")),
            _safe(asset,"page_url"),
            f"{float(_safe(asset,'match_confidence',default='0'))*100:.0f}%",
            _safe(asset,"verification_status"),
        ]
        for ci,val in enumerate(vals): row.cells[ci].text=val
        if ri%2==0:
            for cell in row.cells: _cell_bg(cell,"E8EBF8")

def _compliance_detail(doc, scorecard, brand):
    _heading(doc, "Compliance Status Detail", 1, brand)
    status = scorecard.get("traiga_status","not_assessed")
    _kv(doc,"TRAIGA Status",    STATUS_LABELS.get(status,status))
    _kv(doc,"Compliance Score", f"{_safe(scorecard,'compliance_score')}/100")
    _kv(doc,"Compliance Band",  _safe(scorecard,"band").upper())
    _kv(doc,"Open Violations",  str(scorecard.get("open_violations_count","N/A")))
    _kv(doc,"Min Days Remaining",str(scorecard.get("min_days_remaining","N/A")))
    _kv(doc,"Last Scanned (UTC)",_safe(scorecard,"last_scanned_utc"))
    _kv(doc,"Domain Audited",   _safe(scorecard,"domain"))
    doc.add_paragraph()
    mp = doc.add_paragraph(); mp.add_run("Scoring Methodology").bold=True
    doc.add_paragraph("Compliance scores are computed on a 0-100 scale. A score of 100 indicates full compliance with all assessed TRAIGA requirements. Deductions are applied based on violation severity (high: -30, medium: -15, low: -5) and whether violations remain open past the 60-day cure period (additional -10 per expired violation).")

def _violations_section(doc, violations, brand):
    _heading(doc, "Violations & Cure Period", 1, brand)
    open_v = [v for v in violations if v.get("status") in ("in_cure","expired","non_compliant")]
    if not open_v:
        doc.add_paragraph("No open violations detected as of the last scan.")
        return
    doc.add_paragraph(f"{len(open_v)} open violation(s) identified. Each violation carries a 60-day cure period from the date of first observation per Tex. Bus. & Com. Code §552.053.")
    doc.add_paragraph()
    for i,v in enumerate(open_v,start=1):
        _heading(doc, f"Violation {i}: {_safe(v,'rule_id')} -- {_safe(v,'citation')}", 2, brand)
        sev = v.get("severity","medium")
        sc  = SEVERITY_COLORS.get(sev, SEVERITY_COLORS["medium"])
        _kv(doc,"Severity",      sev.upper(), sc)
        _kv(doc,"Status",        v.get("status","N/A").replace("_"," ").title())
        _kv(doc,"First Observed",_safe(v,"first_observed_utc").split("T")[0])
        cite = _safe(v,"citation")
        _kv(doc,"Citation", cite, url=_citation_url(cite))
        _kv(doc,"Cure Deadline", _safe(v,"cure_deadline_utc").split("T")[0])
        days_raw = v.get("days_remaining")
        try:
            days = int(days_raw) if days_raw not in (None, "", "None", "NaN") else None
        except (ValueError, TypeError):
            days = None
        if days is not None:
            dc = SEVERITY_COLORS["high"] if days<15 else (SEVERITY_COLORS["medium"] if days<30 else None)
            _kv(doc,"Days Remaining",str(days),dc)
        _kv(doc,"Domain",_safe(v,"domain"))
        _kv(doc,"Vendor",_safe(v,"vendor_id"))
        evidence = v.get("evidence",{})
        if isinstance(evidence,dict) and evidence:
            doc.add_paragraph()
            ep = doc.add_paragraph(); ep.add_run("Evidence").bold=True
            _kv(doc,"Page URL",_safe(evidence,"page_url"))
            inds = evidence.get("matched_indicators",[])
            if inds: _kv(doc,"Matched Indicators",", ".join(str(x) for x in inds))
            rem = _safe(evidence,"remediation")
            if rem!="N/A": _kv(doc,"Recommended Remediation",rem)
        doc.add_paragraph()

def _recommendations(doc, violations, brand):
    _heading(doc, "Remediation Recommendations", 1, brand)
    open_v = [v for v in violations if v.get("status") in ("in_cure","expired","non_compliant")]
    if not open_v:
        doc.add_paragraph("No remediation actions are currently required. To maintain compliance, ensure that any future AI system deployments include conspicuous disclosure language per Tex. Bus. & Com. Code §552.051 before going live.")
        return
    doc.add_paragraph("The following actions are recommended in priority order. HIGH severity items should be addressed immediately.")
    doc.add_paragraph()
    seen = set()
    for sev in ("high","medium","low"):
        for v in open_v:
            if v.get("severity")!=sev: continue
            rule_id = _safe(v,"rule_id")
            if rule_id in seen: continue
            seen.add(rule_id)
            evidence = v.get("evidence",{})
            rem = (_safe(evidence,"remediation") if isinstance(evidence,dict) else "N/A")
            p = doc.add_paragraph(style="List Number")
            h = p.add_run(f"[{sev.upper()}] {rule_id}: "); h.bold=True; h.font.color.rgb=SEVERITY_COLORS.get(sev,SEVERITY_COLORS["medium"])
            p.add_run(rem if rem!="N/A" else f"Review and remediate violation of {_safe(v,'citation')}.")
    doc.add_paragraph()
    doc.add_paragraph("All disclosure language should clearly state that the user is interacting with an automated AI system (not a human), be visible before or at the start of each interaction, and be written in plain language accessible to the general public.")

def _statutory_reference(doc, brand):
    _heading(doc, "Statutory Reference -- TRAIGA Key Provisions", 1, brand)
    doc.add_paragraph("The following provisions of Texas HB 149 (89th Legislature), codified as Tex. Bus. & Com. Code Ch. 552, are most relevant to this assessment.")
    doc.add_paragraph()
    for citation,title,text in TRAIGA_PROVISIONS:
        _heading(doc, f"{citation} -- {title}", 2, brand)
        doc.add_paragraph(text)
        # Link directly to the statutory text
        p = doc.add_paragraph()
        p.add_run("View statute: ").font.size = Pt(9)
        _add_hyperlink(p, f"{STATUTE_BASE}#{citation.lstrip('§')}", f"{STATUTE_BASE}#{citation.lstrip('§')}", "1565C0")
        doc.add_paragraph()
    doc.add_paragraph()
    dp = doc.add_paragraph(); dp.add_run("Disclaimer").bold=True
    doc.add_paragraph("This report is generated by the AI Transparency Auditor Platform through automated scanning of publicly accessible web resources. It does not constitute legal advice. Municipalities should consult with legal counsel before making compliance determinations or taking enforcement-related action. Statutory text references are provided for informational purposes and should be verified against the official Texas statutes at statutes.capitol.texas.gov.")

def generate(city, scorecard, violations, output_path, brand):
    brand       = _load_brand(brand if isinstance(brand, str) else None)
    doc         = Document()
    doc_id      = f"TRAIGA-{datetime.now(timezone.utc).strftime('%Y%m%d')}-{str(uuid.uuid4())[:6].upper()}"
    report_date = datetime.now(timezone.utc).strftime("%B %d, %Y")
    _setup_document(doc, brand)
    _cover_page(doc, city, report_date, doc_id, brand)
    _executive_summary(doc, city, scorecard, violations, brand)
    _asset_inventory(doc, scorecard, brand)
    _compliance_detail(doc, scorecard, brand)
    _violations_section(doc, violations, brand)
    _recommendations(doc, violations, brand)
    _statutory_reference(doc, brand)
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="TRAIGA Compliance Report Generator")
    parser.add_argument("--city",         required=True)
    parser.add_argument("--scorecard",    required=True)
    parser.add_argument("--violations",   required=True)
    parser.add_argument("--output",       required=True)
    parser.add_argument("--brand-config", default=None, dest="brand_config")
    args = parser.parse_args()
    try:
        sc = json.loads(args.scorecard)
        vl = json.loads(args.violations)
    except json.JSONDecodeError as e:
        print(f"ERROR: Invalid JSON -- {e}"); sys.exit(1)
    brand = _load_brand(args.brand_config)
    out   = generate(args.city, sc, vl, args.output, brand)
    print(f"Report saved to: {out}")
