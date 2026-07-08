"""
safeharbor.py — TRAIGA safe-harbor readiness (Municipal AI Profile) endpoints.

GET  /api/safeharbor/{city}            readiness checklist + scores (RBAC-scoped read)
POST /api/safeharbor/{city}/attest     attest / clear a control (platform_admin or
                                       agency_admin holding the city)
GET  /api/safeharbor/{city}/statement  NIST AI RMF Alignment Statement (.docx)

Legal basis: Tex. Bus. & Com. Code Sec. 552.105(c)-(e) (HB 149 enrolled; verified
2026-07-07). Generated statements carry a counsel-review disclaimer.
"""
from __future__ import annotations

import os
import tempfile
from datetime import datetime, timezone
from typing import Any, Dict, Optional

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel

from core.access import resolve_principal
from core.auth import get_current_user
from core.dependencies import get_repository
from core.governance_service import GovernanceRepository
from core.safeharbor import build_context, evaluate_profile
from engine.rule_loader import load_schema

router = APIRouter(prefix="/safeharbor", tags=["safeharbor"])


class AttestBody(BaseModel):
    control_id: str
    status: str = "attested"          # "attested" | "open" (clears)
    notes: str = ""


def _module() -> Dict[str, Any]:
    schema = load_schema()
    module = schema.get("Safe_Harbor_Module")
    if not module:
        raise HTTPException(status_code=500, detail="Safe_Harbor_Module missing from schema")
    return module


def _require_read(principal, city: str) -> None:
    if not principal.can_see_city(city):
        raise HTTPException(status_code=403, detail="Not authorized for this city")


def _require_attest(principal, city: str) -> None:
    if principal.is_platform_admin:
        return
    if principal.is_agency_admin and principal.can_see_city(city):
        return
    raise HTTPException(status_code=403,
                        detail="Attestation requires platform or agency administrator rights for this city")


@router.get("/{city}")
def get_readiness(
    city: str,
    repo: GovernanceRepository = Depends(get_repository),
    user: dict = Depends(get_current_user),
):
    principal = resolve_principal(user, repo)
    _require_read(principal, city)
    module = _module()
    ctx = build_context(repo, city)
    attestations = repo.get_safe_harbor(city)
    result = evaluate_profile(module, ctx, attestations)
    result["city"] = city
    return result


@router.post("/{city}/attest")
def attest_control(
    city: str,
    body: AttestBody,
    repo: GovernanceRepository = Depends(get_repository),
    user: dict = Depends(get_current_user),
):
    principal = resolve_principal(user, repo)
    _require_attest(principal, city)
    module = _module()
    valid_ids = {c["control_id"] for c in module.get("controls", [])}
    if body.control_id not in valid_ids:
        raise HTTPException(status_code=400, detail=f"Unknown control_id {body.control_id}")
    if body.status not in ("attested", "open"):
        raise HTTPException(status_code=400, detail="status must be 'attested' or 'open'")

    record = {
        "city": city,
        "control_id": body.control_id,
        "status": body.status,
        "notes": body.notes[:2000],
        "attested_by": user.get("email", "unknown"),
        "attested_utc": datetime.now(timezone.utc).isoformat(),
        "profile_version": module.get("profile_version", "1.0"),
    }
    repo.upsert_safe_harbor(record)
    try:
        repo.append_audit_log(event="safe_harbor_attested", city_count=0, failures=0, details={
            "actor": user.get("email", "unknown"),
            "summary": f"{city}: {body.control_id} -> {body.status}",
            "city": city, "control_id": body.control_id, "status": body.status,
        })
    except Exception as exc:
        print(f"[safeharbor] WARN: audit log failed: {type(exc).__name__}: {exc}")
    return record


@router.get("/{city}/statement")
def alignment_statement(
    city: str,
    repo: GovernanceRepository = Depends(get_repository),
    user: dict = Depends(get_current_user),
):
    """Generate the NIST AI RMF Alignment Statement (.docx) for a city."""
    principal = resolve_principal(user, repo)
    _require_read(principal, city)
    module = _module()
    ctx = build_context(repo, city)
    attestations = repo.get_safe_harbor(city)
    result = evaluate_profile(module, ctx, attestations)

    path = os.path.join(tempfile.mkdtemp(prefix="safeharbor_"),
                        f"{city.replace(' ', '_')}_Alignment_Statement.docx")
    try:
        _build_statement_docx(path, city, module, result, ctx)
    except Exception as exc:
        raise HTTPException(status_code=500,
                            detail=f"Statement generation failed: {type(exc).__name__}: {exc}")
    try:
        repo.append_audit_log(event="alignment_statement_generated", city_count=0, failures=0, details={
            "actor": user.get("email", "unknown"),
            "summary": f"Alignment Statement generated for {city} "
                       f"({result['overall']['satisfied']}/{result['overall']['total']} controls)",
            "city": city,
        })
    except Exception:
        pass
    return FileResponse(
        path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=os.path.basename(path),
    )


# ── DOCX builder ──────────────────────────────────────────────────────────────

_STATUS_LABEL = {"satisfied": "Satisfied", "failing": "Not satisfied", "open": "Attestation required"}


def _build_statement_docx(path: str, city: str, module: Dict[str, Any],
                          result: Dict[str, Any], ctx: Dict[str, Any]) -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    NAVY = RGBColor(0x1A, 0x3C, 0x6E)
    doc = Document()

    def h(text, level=1):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            run.font.color.rgb = NAVY
        return p

    title = doc.add_heading(f"NIST AI RMF Alignment Statement — {city}", level=0)
    for run in title.runs:
        run.font.color.rgb = NAVY

    now = datetime.now(timezone.utc)
    meta = doc.add_paragraph()
    meta.add_run(f"{module.get('profile_name')} v{module.get('profile_version')}  ·  "
                 f"Generated {now.strftime('%B %d, %Y')} (UTC)  ·  TRAIGA Auditor").italic = True

    h("Purpose and Legal Basis", 1)
    lb = module.get("legal_basis", {})
    doc.add_paragraph(
        f"This statement documents {city}'s alignment with the {module.get('profile_name')} — a municipal "
        "profile of the NIST AI Risk Management Framework (AI RMF 1.0) and the NIST Artificial Intelligence "
        "Risk Management Framework: Generative Artificial Intelligence Profile (NIST AI 600-1). "
        f"It is maintained to support the defenses available under {lb.get('citation', 'Tex. Bus. & Com. Code Sec. 552.105')}: "
        "the rebuttable presumption of reasonable care (552.105(c)) and the defense for violations discovered "
        "through testing or an internal review process substantially complying with a nationally recognized AI "
        "risk management framework (552.105(e)).")

    h("Internal Review Process and Testing Program", 1)
    row = ctx.get("scorecard_row") or {}
    violations = ctx.get("violations") or []
    cured = [v for v in violations if v.get("status") == "cured"]
    open_v = [v for v in violations if v.get("status") in ("open", "in_cure")]
    doc.add_paragraph(
        f"{city}'s program operates continuous, automated external testing of its public digital services for "
        "artificial intelligence systems and statutory disclosure compliance (Sec. 552.105(e)(B) testing), an "
        "AI use-case inventory with human attestation of purpose and ownership, and statutory 60-day cure "
        "tracking for every finding (Sec. 552.105(e)(D) internal review process). "
        f"Most recent assessment: {row.get('last_scanned_utc') or 'not yet assessed'}. "
        f"Findings to date: {len(violations)} total; {len(cured)} remediated; {len(open_v)} in active cure.")

    h("Readiness Summary", 1)
    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, txt in enumerate(("Function", "Controls satisfied", "Total", "Percent")):
        hdr[i].text = txt
    for fn_name in ("govern", "map", "measure", "manage"):
        s = result["scores"].get(fn_name, {})
        r = t.add_row().cells
        r[0].text = fn_name.upper()
        r[1].text = str(s.get("satisfied", 0))
        r[2].text = str(s.get("total", 0))
        r[3].text = f"{int(s.get('pct', 0) * 100)}%"
    o = result["overall"]
    p = doc.add_paragraph()
    p.add_run(f"Overall: {o['satisfied']} of {o['total']} controls satisfied "
              f"({int(o['pct'] * 100)}%) — readiness band: {result['band'].upper()}.").bold = True

    h("Control Detail and Evidence", 1)
    ct = doc.add_table(rows=1, cols=5)
    ct.style = "Light Grid Accent 1"
    hdr = ct.rows[0].cells
    for i, txt in enumerate(("Control", "NIST reference", "Status", "Basis", "Evidence / attestation")):
        hdr[i].text = txt
    for c in result["controls"]:
        r = ct.add_row().cells
        r[0].text = f"{c['control_id']} — {c['title']}"
        r[1].text = c.get("nist_ref", "")
        r[2].text = _STATUS_LABEL.get(c["status"], c["status"])
        r[3].text = ("Machine-verified from platform data" if c["basis"] == "machine"
                     else "Human attestation" if c["basis"] == "attested"
                     else "Awaiting attestation")
        att = c.get("attestation")
        if att and att.get("status") == "attested":
            r[4].text = (f"Attested by {att.get('attested_by')} on "
                         f"{str(att.get('attested_utc'))[:10]}. {att.get('notes') or ''}").strip()
        elif c["basis"] == "machine":
            r[4].text = "Derived from scan history, inventory, violation and audit-log records held by the platform."
        else:
            r[4].text = c.get("attest_hint", "")

    h("Limitations and Counsel Review", 1)
    doc.add_paragraph(
        "This statement is generated from platform records and human attestations as of the date above. "
        "The presumption under Sec. 552.105(c) is rebuttable, and the defenses under Sec. 552.105(e) depend on "
        "facts and circumstances. This document is an engineering artifact, not legal advice; review by "
        "counsel is required before reliance or production to any regulator.")

    doc.add_paragraph("\n\nAcknowledged for the city:")
    sig = doc.add_paragraph("\n____________________________\nName / Title / Date")
    sig.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for section in doc.sections:
        section.footer.paragraphs[0].text = (
            f"{city} — NIST AI RMF Alignment Statement — generated {now.strftime('%Y-%m-%d')} — counsel review required")
    doc.save(path)
