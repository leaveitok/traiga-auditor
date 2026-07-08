"""
cid.py — AG Civil Investigative Demand endpoints.

GET /api/cid/{city}/readiness        per-asset, per-item 552.103(b) readiness
GET /api/cid/{city}/package          AG Response Package (.docx)
GET /api/cid/{city}/cure-statement   552.104(b)(2) Cure Statement (.docx)

RBAC: reads scoped by can_see_city (same model as safeharbor). Field editing
happens through the inventory PATCH route, which enforces write rights.
"""
from __future__ import annotations

import os
import tempfile
from datetime import datetime, timezone
from typing import Any, Dict, List

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse

from core.access import resolve_principal
from core.auth import get_current_user
from core.cid import CID_CITATION, CURE_CITATION, evaluate_city
from core.dependencies import get_repository
from core.governance_service import GovernanceRepository
from core.safeharbor import build_context

router = APIRouter(prefix="/cid", tags=["cid"])


def _require_read(principal, city: str) -> None:
    if not principal.can_see_city(city):
        raise HTTPException(status_code=403, detail="Not authorized for this city")


def _readiness(repo: GovernanceRepository, city: str) -> Dict[str, Any]:
    ctx = build_context(repo, city)
    assets = repo.get_ai_assets(city=city)
    result = evaluate_city(assets, ctx)
    result["city"] = city
    return result


@router.get("/{city}/readiness")
def cid_readiness(
    city: str,
    repo: GovernanceRepository = Depends(get_repository),
    user: dict = Depends(get_current_user),
):
    principal = resolve_principal(user, repo)
    _require_read(principal, city)
    return _readiness(repo, city)


@router.get("/{city}/package")
def cid_package(
    city: str,
    repo: GovernanceRepository = Depends(get_repository),
    user: dict = Depends(get_current_user),
):
    """AG Response Package: sections mirror 552.103(b)(1)-(8) per AI system."""
    principal = resolve_principal(user, repo)
    _require_read(principal, city)
    readiness = _readiness(repo, city)
    ctx = build_context(repo, city)
    violations = repo.get_violations(city=city)

    path = os.path.join(tempfile.mkdtemp(prefix="cid_"),
                        f"{city.replace(' ', '_')}_AG_Response_Package.docx")
    try:
        _build_package_docx(path, city, readiness, ctx, violations)
    except Exception as exc:
        raise HTTPException(status_code=500,
                            detail=f"Package generation failed: {type(exc).__name__}: {exc}")
    _log(repo, user, "cid_package_generated",
         f"AG Response Package generated for {city} "
         f"({readiness['ready_count']}/{readiness['asset_count']} assets fully answerable)", city)
    return FileResponse(path, filename=os.path.basename(path),
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@router.get("/{city}/cure-statement")
def cure_statement(
    city: str,
    repo: GovernanceRepository = Depends(get_repository),
    user: dict = Depends(get_current_user),
):
    """552.104(b)(2) written cure statement covering the city's cured violations."""
    principal = resolve_principal(user, repo)
    _require_read(principal, city)
    violations = repo.get_violations(city=city)
    cured = [v for v in violations if v.get("status") == "cured"]
    if not cured:
        raise HTTPException(status_code=404,
                            detail="No cured violations on record — a cure statement requires at least one.")
    ctx = build_context(repo, city)
    path = os.path.join(tempfile.mkdtemp(prefix="cure_"),
                        f"{city.replace(' ', '_')}_Cure_Statement.docx")
    try:
        _build_cure_docx(path, city, cured, ctx)
    except Exception as exc:
        raise HTTPException(status_code=500,
                            detail=f"Cure statement generation failed: {type(exc).__name__}: {exc}")
    _log(repo, user, "cure_statement_generated",
         f"Cure Statement generated for {city} ({len(cured)} cured violations)", city)
    return FileResponse(path, filename=os.path.basename(path),
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


def _log(repo, user, event, summary, city):
    try:
        repo.append_audit_log(event=event, city_count=0, failures=0, details={
            "actor": user.get("email", "unknown"), "summary": summary, "city": city})
    except Exception as exc:
        print(f"[cid] WARN: audit log failed: {type(exc).__name__}: {exc}")


# ── DOCX builders ─────────────────────────────────────────────────────────────

def _docx_common():
    from docx import Document
    from docx.shared import RGBColor
    doc = Document()
    return doc, RGBColor(0x1A, 0x3C, 0x6E)


def _heading(doc, navy, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = navy
    return p


_SOURCE_LABEL = {
    "machine": "Platform-derived (automated observation)",
    "attested": "Documented by city staff",
    "vendor_referred": "Vendor-operated; referred to vendor",
    "composite": "Composed from platform monitoring records",
}


def _build_package_docx(path: str, city: str, readiness: Dict[str, Any],
                        ctx: Dict[str, Any], violations: List[Dict[str, Any]]) -> None:
    doc, navy = _docx_common()
    now = datetime.now(timezone.utc)

    title = doc.add_heading(f"Civil Investigative Demand Response Package — {city}", level=0)
    for run in title.runs:
        run.font.color.rgb = navy
    doc.add_paragraph(
        f"Prepared {now.strftime('%B %d, %Y')} (UTC) · TRAIGA Auditor · Structured to {CID_CITATION}. "
        "DRAFT for counsel review — not for production to any regulator without attorney approval."
    ).runs[0].italic = True

    _heading(doc, navy, "How to Use This Package", 1)
    doc.add_paragraph(
        "Sections below answer, per artificial intelligence system, the eight categories of information the "
        "Office of the Attorney General may request under Sec. 552.103(b). Each answer is labeled with its "
        "evidence source. Items marked GAP require completion in the AI inventory before this package is "
        "production-ready. All platform-derived answers regenerate from live records each time this document "
        "is produced.")

    if not readiness["assets"]:
        doc.add_paragraph("No active AI systems are on record for this city. The registry, scan history, and "
                          "monitoring records in the appendix support a response that no covered systems are deployed.")
    for a in readiness["assets"]:
        _heading(doc, navy, f"AI System: {a['display_name']}", 1)
        if a["vendor_operated"]:
            doc.add_paragraph("Classification: vendor-operated system deployed by the city.")
        t = doc.add_table(rows=1, cols=3)
        t.style = "Light Grid Accent 1"
        hdr = t.rows[0].cells
        for i, txt in enumerate(("552.103(b) item", "Response", "Source")):
            hdr[i].text = txt
        for item in a["items"]:
            r = t.add_row().cells
            r[0].text = f"({item['item']}) {item['title']}"
            r[1].text = ("[GAP — complete before production] " if item["status"] == "gap" else "") + item["text"]
            r[2].text = _SOURCE_LABEL.get(item["source"], "—") if item["source"] else "MISSING"
        summary = doc.add_paragraph()
        summary.add_run(
            f"Readiness: {a['answered']} of {a['total']} items answerable."
            + ("" if a["ready"] else f" Open gaps: {', '.join(a['gaps'])}.")).bold = True

    _heading(doc, navy, "Appendix A — Compliance Posture", 1)
    row = ctx.get("scorecard_row") or {}
    doc.add_paragraph(
        f"Current status: {row.get('traiga_status', 'not assessed')} · last assessed "
        f"{str(row.get('last_scanned_utc') or 'n/a')[:10]} · open violations: "
        f"{sum(1 for v in violations if v.get('status') in ('open', 'in_cure'))} · cured: "
        f"{sum(1 for v in violations if v.get('status') == 'cured')}.")
    _heading(doc, navy, "Appendix B — Violation and Remediation History", 1)
    if violations:
        vt = doc.add_table(rows=1, cols=5)
        vt.style = "Light Grid Accent 1"
        hdr = vt.rows[0].cells
        for i, txt in enumerate(("Rule", "Severity", "Status", "First observed", "Cure deadline")):
            hdr[i].text = txt
        for v in violations:
            r = vt.add_row().cells
            r[0].text = str(v.get("rule_id", ""))
            r[1].text = str(v.get("severity", ""))
            r[2].text = str(v.get("status", ""))
            r[3].text = str(v.get("first_observed_utc", ""))[:10]
            r[4].text = str(v.get("cure_deadline_utc", ""))[:10]
    else:
        doc.add_paragraph("No violations on record.")

    _heading(doc, navy, "Limitations", 1)
    doc.add_paragraph(
        "Engineering artifact generated from platform records; findings are candidate compliance signals. "
        "Counsel review is required before reliance or production. The city's NIST AI RMF Alignment Statement "
        "(generated separately) documents the internal review process supporting Sec. 552.105 defenses.")
    for section in doc.sections:
        section.footer.paragraphs[0].text = (
            f"{city} — CID Response Package (DRAFT, counsel review required) — {now.strftime('%Y-%m-%d')}")
    doc.save(path)


def _build_cure_docx(path: str, city: str, cured: List[Dict[str, Any]],
                     ctx: Dict[str, Any]) -> None:
    doc, navy = _docx_common()
    now = datetime.now(timezone.utc)

    title = doc.add_heading(f"Written Statement of Cure — {city}", level=0)
    for run in title.runs:
        run.font.color.rgb = navy
    doc.add_paragraph(
        f"Prepared {now.strftime('%B %d, %Y')} (UTC) · Structured to {CURE_CITATION}. "
        "DRAFT for counsel review before submission to the Office of the Attorney General."
    ).runs[0].italic = True

    _heading(doc, navy, "(i) Statement of Cure", 1)
    doc.add_paragraph(
        f"{city} states that the violations identified below have been cured. Cure was verified by automated "
        "re-assessment of the affected public digital services; verification records are reproduced in (ii).")

    _heading(doc, navy, "(ii) Supporting Documentation — Manner of Cure", 1)
    t = doc.add_table(rows=1, cols=5)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, txt in enumerate(("Rule / citation", "Severity", "First observed", "Verified cured", "Evidence")):
        hdr[i].text = txt
    for v in cured:
        r = t.add_row().cells
        r[0].text = f"{v.get('rule_id', '')} — {v.get('citation', '')}"
        r[1].text = str(v.get("severity", ""))
        r[2].text = str(v.get("first_observed_utc", ""))[:10]
        r[3].text = str(v.get("last_observed_utc", "") or v.get("cured_utc", ""))[:10]
        r[4].text = ("Automated re-scan no longer observes the violating condition; scan evidence retained "
                     "in the platform's violation record.")

    _heading(doc, navy, "(iii) Internal Policy Changes to Prevent Recurrence", 1)
    row = ctx.get("scorecard_row") or {}
    doc.add_paragraph(
        "The city maintains: continuous automated external testing of its public digital services "
        f"(most recent assessment {str(row.get('last_scanned_utc') or 'n/a')[:10]}); an AI use-case inventory "
        "with owner attestation; a Municipal AI Profile (NIST AI RMF) governance checklist with recorded "
        "attestations; and statutory 60-day cure tracking for all findings. [Counsel: attach the adopted AI "
        "use policy and describe any additional policy changes specific to these violations.]")

    _heading(doc, navy, "Signature", 1)
    doc.add_paragraph("\nFor the city:\n\n____________________________\nName / Title / Date")
    for section in doc.sections:
        section.footer.paragraphs[0].text = (
            f"{city} — Written Statement of Cure (DRAFT, counsel review required) — {now.strftime('%Y-%m-%d')}")
    doc.save(path)
